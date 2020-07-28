import pyttsx3
import datetime
import speech_recognition as sr
import wikipedia
import webbrowser
import os
import googlesearch
import docx
import smtplib
import subprocess
import requests
from bs4 import BeautifulSoup



#Initializing the engine
engine = pyttsx3.init('sapi5')
voice = engine.getProperty('voices')
engine.setProperty('voice', voice[0].id)


#Speaks the computer
def speak(audio):
    engine.say(audio)
    engine.runAndWait()


mail_person = {
    'dad': 'emailofdad@gmail.com',
    'mom': 'emailofmom@gmail.com'
}



def speaker():
    #Takes the voice command from the user and returns it in a string format
    rec = sr.Recognizer()
    with sr.Microphone() as mic:
        print("Listening...")
        rec.pause_threshold = 1 #Here if, the user stops for 1 sec, still the compiler will not end.
        audio = rec.listen(mic)
    try:
        print("Recognizing...")
        query = rec.recognize_google(audio,language='en-in') #Searches at the google for recognisation
        print(f"User said: {query}\n") #Prints the query

    except Exception as e:
        print("Sorry sir, I could not get it")
        speak("Sorry sir, I could not get it")
        return "None"
    return query



def wish_on_time():
    # Wishes the user based on the time
    hou = int(datetime.datetime.now().hour)
    if hou >= 5 and hou<=12:
        speak("Good Morning!")
    elif hou>=12 and hou<=16:
        speak("Good Afternoon!")
    elif hou>=17 and hou<=20:
        speak("Good Evening!")
    else:
        speak("Good Night!")
    speak("I am JARVIS sir, How may I help you?")

def send_mail(recepient,content):
    server = smtplib.SMTP('smtp.gmail.com',587)
    server.ehlo()
    server.starttls()
    server.login('myemail@gmail.com','my-password')    #Write your email id and password here
    server.sendmail('myemail@gmail.com',recepient,content)
    server.close()



if __name__ == '__main__':
    wish_on_time()
    while True:
        command = speaker().lower()
        if 'how are you' in command:
            speak('I am fine sir. Hope you are also well!')

        elif 'what can you do' in command:
            speak("Sir, I can search any topic on wikipedia, google and can give you the exact result, I can"
                  " open youtube, google. I can give you the perfect location of any places along with any route "
                  ".I can tell you the time, date. I can make you feel relax by playing music. I can open"
                  " calculator, dev C++. I can read full articles in newspaper for you, write for you and send email"
                  " for you")

        elif 'wikipedia' in command:  #Searches at wikipedia if the user uses the word wikipedia
            speak("Searching results for you in wikipedia...")
            command = command.replace('wikipedia','').replace("search","").replace('about','')
            result = wikipedia.summary(command, sentences=3)   #Takes the first three statements from the wikipedia as result
            print(result)
            speak("According to wikipedia")
            speak(result)

        elif 'search google about' in command:
            command = command.replace('search google about', '')
            speak("Opening search results about" + command)
            for i in googlesearch.search(command,tld='co.in',lang='en',num=10,stop=1,pause=2):  #forms the url
                print(i)
                webbrowser.open(i)     #Opens the page

        elif 'google map' in command:
            command = command.replace('google map','')
            command = command.replace('search','')
            command = command.replace('about','').replace('in','')
            if 'route' in command or 'path' in command:
                speak('Sir! What is your starting location?')
                start = speaker()
                speak('Sir! What is the destination?')
                end = speaker()
                speak('Opening the route for you, sir!')
                webbrowser.open('https://www.google.co.in/maps/dir/' + start + '/' + end)
            else:
                speak('Opening search results for you, sir!')
                webbrowser.open('https://www.google.com/maps/search/' + command)

        elif 'open google' in command:
            speak('opening google for you, sir')
            webbrowser.open('https://www.google.com')   #Opens google

        elif 'open youtube' in command:
            speak('Opening youtube for you, sir')
            webbrowser.open('https://www.youtube.com')   #Opens youtube

        elif 'play music' in command:
            music_directory = 'Your music directory'  #Music directory
            speak('Enjoy sir!')
            songs = os.listdir(music_directory)
            os.startfile(os.path.join(music_directory,songs[0]))   #Starts the music file by joining the path i.e. the directory with the songs

        elif 'the time' in command:
            speak("Sir, the time is")
            speak(datetime.datetime.now().strftime("%H:%M:%S"))     #Speaks the time

        elif 'the date' in command:
            speak('Sir, the date is')
            speak(datetime.datetime.now().date())    #Speaks the date

        elif 'code' in command:
            speak('Opening Dev C++')
            dev_dir = "Your dev c++ directory"
            os.startfile(dev_dir)
        elif 'calculator' in command:
            speak('Starting calculator for you, sir')
            subprocess.Popen('Your calculator directory')

        elif 'start writing' in command:
            speak("What is your document name,sir!")
            save_name = speaker()
            while save_name == 'None':
                save_name = speaker()
            speak('Speak sir!')
            mydoc = docx.Document()
            while True:
                speech = speaker()
                command = mydoc.add_paragraph(speech)
                mydoc.save("The document directory"+save_name+".docx")
                speak('You want to stop or add heading sir?')
                attribute = speaker()
                while attribute == "None":
                    speak('Sorry sir! Will I add heading or stop the document?')
                    attribute = speaker()
                if 'stop' in attribute:
                    speak('Your document has been saved successfully, sir')
                    break
                elif 'add heading' in attribute:
                    speak('Heading please sir:')
                    mydoc.add_heading(speaker(),1)
                else:
                    speak("Ok! continue speaking sir.")
                    continue

        elif 'send email' in command:
            try:
                speak("Whom do you want to say sir?")
                person = speaker().lower()
                recipient = ''
                while person == 'none' or not(person in mail_person):
                    person = speaker().lower()
                    if 'me' in person:
                        recipient = mail_person['mom']
                        break
                    elif 'dad' in person:
                        recipient = mail_person['dad']
                        break
                speak("What do you want to speak sir?")
                content = speaker()
                while content == 'None':
                    speak("Sir, I cannot recognise, Please speak again sir.")
                    content = speaker()
                send_mail(recipient,content)
                speak("Email has been sent successfully")
            except Exception as e:
                print(e)
                speak("Sorry sir, i think there is an error in sending mail.")

        elif 'newspaper' in command:
            speak("What is your news topic, sir!")
            search = speaker()
            search = search.replace("seach","").replace("about","").replace("news","")
            url = "https://www.thestatesman.com/?s="+search+"&search_btn_inner=Search"
            speak("Opening news about"+search+", It may take some time.")

            c = requests.get(url)
            r = c.content
            soup = BeautifulSoup(r, "html.parser")
            all = soup.find_all("div", {"class": "content-block"})
            speak("There are "+str(len(all))+" news on this toic")
            for j in range(len(all)):
                detail = all[j].find("a")
                inner_url = detail.attrs['href']
                details = detail.text
                speak("News number "+str(j+1))
                speak(details)
                speak("Sir,do you want to read the full article on this topic?")
                ans = speaker()
                if 'yes' in ans or 'obviously' in ans or 'read' in ans or 'yeah' in ans:
                    c1 = requests.get(inner_url)
                    r1 = c1.content
                    SOUP = BeautifulSoup(r1, "html.parser")

                    news_para = SOUP.find_all("p")
                    for i in range(0, len(news_para)):
                        news = news_para[i].text
                        print(news)
                        speak(news)
                elif 'stop' in ans or 'break' in ans:
                    break
                else:
                    continue


        elif 'weather'in command:
            speak("Tell me the name of the place, sir!")
            town = speaker()
            url = "http://api.openweathermap.org/data/2.5/weather?appid=YOUR-API-KEY&q=" + town
            details = requests.get(url).json()

            temp = str(round(details['main']['temp'] - 273, 2))
            feel = str(round(details['main']['feels_like'] - 273, 2))
            min_temp = str(round(details['main']['temp_min'] - 273, 2))
            max_temp = str(round(details['main']['temp_max'] - 273, 2))
            humidity = str(round(details['main']['humidity'], 2))
            weather = details['weather'][0]['main']

            detailed_report = "Temperature is " + temp + " degree centigrade. Todays maximum temperature is " +\
                              max_temp + " degree centigrade and minimum temperature is " + min_temp \
                              + " degree centigrade. Due to " + weather + " it feels like " \
                              + feel + " degree centigrade. Humidity is " + humidity + " percent"
            print(detailed_report)
            speak(detailed_report)


        elif 'stop' in command:       #Exits the process
            speak("Goodbye sir! Have a nice day")
            break
